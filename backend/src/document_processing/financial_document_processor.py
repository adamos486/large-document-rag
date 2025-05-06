import os
import re
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional, Union

from ..config.config import settings
import logging
import json
import datetime
import hashlib

# Document processing libraries
import pytesseract
from pdf2image import convert_from_path
from pypdf import PdfReader
import docx
import pandas as pd
import numpy as np
from unstructured.partition.auto import partition

# NLP for entity extraction
import spacy
from transformers import pipeline
from transformers_interpret import SequenceClassificationExplainer

# Import base classes
from .base_processor import BaseDocumentProcessor, Document, DocumentChunk

# Set up logging
logger = logging.getLogger(__name__)

class FinancialDocumentProcessor(BaseDocumentProcessor):
    """
    Processor for financial documents used in M&A due diligence.
    Handles various document types including PDFs, Word documents, Excel spreadsheets, etc.
    """
    
    def __init__(self, use_advanced_nlp: bool = True, enable_ocr: bool = True):
        """
        Initialize the financial document processor.
        
        Args:
            use_advanced_nlp: Whether to use advanced NLP for entity extraction.
            enable_ocr: Whether to use OCR for scanned documents.
        """
        self.use_advanced_nlp = use_advanced_nlp
        self.enable_ocr = enable_ocr
        
        # Initialize NLP models for entity extraction if advanced NLP is enabled
        if self.use_advanced_nlp:
            try:
                # Load spaCy model for named entity recognition
                self.nlp = spacy.load("en_core_web_lg")
                
                # Initialize financial entity extraction model using BERT-based models
                # We'll use a general NER model as a replacement for finbert
                self.financial_ner = pipeline(
                    "token-classification", 
                    model="dbmdz/bert-large-cased-finetuned-conll03-english", 
                    aggregation_strategy="simple"
                )
                
                # For financial classification (sentiment/topics)
                try:
                    self.financial_classifier = pipeline(
                        "text-classification",
                        model="yiyanghkust/finbert-tone",
                        return_all_scores=True
                    )
                except Exception as e:
                    logger.warning(f"Could not load financial classifier: {e}")
                    self.financial_classifier = None
            except Exception as e:
                logger.warning(f"Failed to load NLP models: {e}. Falling back to basic processing.")
                self.use_advanced_nlp = False
    
    def load_document(self, file_path: Path) -> Document:
        """
        Load a document from a file path.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            Loaded Document object.
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Generate a document ID based on filename and modification time
        doc_id = self._generate_doc_id(file_path)
        
        # Extract content based on file type
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            content = self._extract_pdf_content(file_path)
        elif file_ext == '.docx':
            content = self._extract_docx_content(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            content = self._extract_excel_content(file_path)
        elif file_ext in ['.csv', '.tsv']:
            content = self._extract_csv_content(file_path)
        elif file_ext == '.txt':
            content = self._extract_text_content(file_path)
        elif file_ext in ['.ppt', '.pptx']:
            content = self._extract_presentation_content(file_path)
        else:
            # Use unstructured library for generic extraction
            try:
                elements = partition(str(file_path))
                content = "\n".join([str(element) for element in elements])
            except Exception as e:
                logger.error(f"Failed to extract content using unstructured: {e}")
                content = f"Unsupported file format: {file_ext}"
        
        # Create basic metadata
        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "file_type": file_ext,
            "doc_id": doc_id,
            "extraction_date": datetime.datetime.now().isoformat(),
        }
        
        return Document(content=content, metadata=metadata, doc_id=doc_id, source=str(file_path))
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate a unique document ID based on file path and modification time."""
        mod_time = os.path.getmtime(file_path)
        unique_str = f"{file_path}_{mod_time}"
        return hashlib.md5(unique_str.encode()).hexdigest()
    
    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract text content from a PDF file."""
        try:
            reader = PdfReader(file_path)
            text = ""
            
            # First try to extract text directly
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            
            # If no text was extracted and OCR is enabled, use OCR
            if not text.strip() and self.enable_ocr:
                logger.info(f"No text extracted directly from PDF, trying OCR: {file_path}")
                images = convert_from_path(file_path)
                
                for i, image in enumerate(images):
                    page_text = pytesseract.image_to_string(image)
                    text += f"Page {i+1}:\n{page_text}\n\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF content: {e}")
            return f"Error extracting content: {str(e)}"
    
    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract text content from a Word document."""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text += " | ".join(row_text) + "\n"
                text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting Word document content: {e}")
            return f"Error extracting content: {str(e)}"
    
    def _extract_excel_content(self, file_path: Path) -> str:
        """Extract content from an Excel spreadsheet."""
        try:
            # Read all sheets
            xlsx = pd.ExcelFile(file_path)
            all_sheets = {}
            
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name)
                all_sheets[sheet_name] = df
            
            # Convert to a structured text representation
            text = ""
            for sheet_name, df in all_sheets.items():
                text += f"\n--- Sheet: {sheet_name} ---\n"
                text += df.to_string(index=False) + "\n\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting Excel content: {e}")
            return f"Error extracting content: {str(e)}"
    
    def _extract_csv_content(self, file_path: Path) -> str:
        """Extract content from a CSV file."""
        try:
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"Error extracting CSV content: {e}")
            return f"Error extracting content: {str(e)}"
    
    def _extract_text_content(self, file_path: Path) -> str:
        """Extract content from a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting text content: {e}")
            return f"Error extracting content: {str(e)}"
    
    def _extract_presentation_content(self, file_path: Path) -> str:
        """Extract content from a PowerPoint presentation."""
        try:
            # Use unstructured for PowerPoint
            elements = partition(str(file_path))
            return "\n\n".join([str(element) for element in elements])
        except Exception as e:
            logger.error(f"Error extracting presentation content: {e}")
            return f"Error extracting content: {str(e)}"
    
    @property
    def chunking_strategy(self):
        """
        Get the current chunking strategy instance.
        
        Returns:
            A chunker instance appropriate for the current document type.
        """
        from .chunking import ChunkerFactory
        
        # Get chunking parameters from settings
        chunk_size = settings.CHUNK_SIZE
        chunk_overlap = settings.CHUNK_OVERLAP
        chunk_strategy = settings.FINANCIAL_CHUNK_STRATEGY
        
        # Create a chunker with the configured parameters
        return ChunkerFactory.create_chunker(
            document_type="financial",
            statement_type=None,  # Will be determined by the chunker
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            respect_financial_boundaries=True
        )

    def chunk_document(self, document: Document) -> List[DocumentChunk]:
        """
        Split a document into chunks optimized for financial document retrieval.
        
        Args:
            document: The document to chunk.
            
        Returns:
            List of document chunks.
        """
        # First, determine document category and type
        metadata = document.metadata or {}
        doc_id = document.doc_id or str(uuid.uuid4())
        file_type = metadata.get("file_type", "")
        
        # Detect document and statement type if possible
        statement_type = metadata.get("statement_type", None)
        document_category = metadata.get("document_category", "financial")
        
        # Add document ID if not present
        if "doc_id" not in metadata:
            metadata["doc_id"] = doc_id
        
        # Different chunking strategies based on document type
        chunks = []
        
        try:
            from .chunking import ChunkerFactory
            
            # For tabular documents (Excel, CSV)
            if file_type in ['.xlsx', '.xls', '.csv', '.tsv']:
                # Special handling for tabular data
                chunks = self._chunk_tabular_document(document.content, metadata, doc_id)
            else:
                # Use our specialized financial chunkers for text documents
                chunker = ChunkerFactory.create_chunker(
                    document_type=document_category,
                    statement_type=statement_type,
                    chunk_size=settings.CHUNK_SIZE,
                    chunk_overlap=settings.CHUNK_OVERLAP,
                    respect_financial_boundaries=True
                )
                
                # Use the chunker to generate chunks
                chunks = chunker.chunk(document)
                
                # If no chunks were generated, fall back to basic chunking
                if not chunks:
                    logger.warning(f"Financial chunker produced no chunks, falling back to basic chunker")
                    chunks = self._basic_chunking(document.content, metadata, doc_id)
        except Exception as e:
            logger.warning(f"Error in specialized chunking: {str(e)}. Falling back to basic chunking.")
            # Fallback to basic chunking if something goes wrong
            chunks = self._basic_chunking(document.content, metadata, doc_id)
        
        # Ensure all chunks have essential metadata
        self._enrich_chunks_with_metadata(chunks, metadata, doc_id)
        
        return chunks

    def _enrich_chunks_with_metadata(self, chunks: List[DocumentChunk], metadata: Dict[str, Any], doc_id: str):
        """
        Enrich chunks with document metadata and financial entity extraction.
        
        Args:
            chunks: List of document chunks to enrich
            metadata: Document metadata
            doc_id: Document identifier
        """
        for chunk in chunks:
            # Ensure core metadata is present# Convert any lists in metadata to strings to ensure compatibility with ChromaDB
            for key, value in chunk.metadata.items():
                if isinstance(value, list):
                    chunk.metadata[key] = ', '.join(value)# Ensure core metadata is present and convert complex types to strings for ChromaDB compatibility
                    for key, value in chunk.metadata.items():
                        if isinstance(value, list):
                            chunk.metadata[key] = ', '.join(map(str, value))
                        elif isinstance(value, dict):
                            # For dictionaries, convert to a formatted string
                            formatted_items = [f"{k}: {v}" for k, v in value.items()]
                            chunk.metadata[key] = '; '.join(formatted_items)
                        elif not isinstance(value, (str, int, float, bool)):
                            # Convert any other non-primitive types to string
                            chunk.metadata[key] = str(value)
            if "doc_id" not in chunk.metadata:
                chunk.metadata["doc_id"] = doc_id
            
            # Add document metadata if not already present
            for key, value in metadata.items():
                if key not in chunk.metadata:
                    chunk.metadata[key] = value
            
            # Always add/update these fields
            chunk.metadata.update({
                "source": metadata.get("source", chunk.metadata.get("source", "")),
                "filename": metadata.get("filename", chunk.metadata.get("filename", "")),
                "file_type": metadata.get("file_type", chunk.metadata.get("file_type", "")),
                "chunk_count": len(chunks)
            })
            
            # Extract financial entities if not already done and advanced NLP is enabled
            if self.use_advanced_nlp and "financial_entities" not in chunk.metadata:
                financial_entities = self._extract_financial_entities(chunk.content)
                if financial_entities:
                    chunk.metadata["financial_entities"] = financial_entities
    
    def _chunk_text_document(self, content: str, metadata: Dict[str, Any], doc_id: str) -> List[DocumentChunk]:
        """
        Chunk a text document using semantic boundaries.
        
        This method attempts to preserve semantic units like paragraphs, sections, etc.
        """
        chunks = []
        
        # Split by sections using various separators
        sections = re.split(r'(\n\s*\n|\r\n\s*\r\n|={3,}|-{3,}|\*{3,})', content)
        
        # Combine section separators with content
        combined_sections = []
        for i in range(0, len(sections)-1, 2):
            if i+1 < len(sections):
                combined_sections.append(sections[i] + sections[i+1])
            else:
                combined_sections.append(sections[i])
        
        if not combined_sections and sections:
            combined_sections = sections
        
        # Further split large sections
        max_chunk_size = 1000  # characters
        current_chunk = ""
        chunk_id = 0
        
        for section in combined_sections:
            if not section.strip():
                continue
                
            if len(current_chunk) + len(section) < max_chunk_size:
                current_chunk += section
            else:
                if current_chunk:
                    chunk_id += 1
                    chunk_metadata = {
                        "chunk_id": f"{doc_id}_{chunk_id}",
                        "chunk_index": chunk_id,
                    }
                    
                    # Extract financial entities if advanced NLP is enabled
                    if self.use_advanced_nlp:
                        financial_entities = self._extract_financial_entities(current_chunk)
                        if financial_entities:
                            chunk_metadata["financial_entities"] = financial_entities
                    
                    chunks.append(DocumentChunk(
                        content=current_chunk,
                        metadata=chunk_metadata,
                        chunk_id=f"{doc_id}_{chunk_id}"
                    ))
                    current_chunk = section
                else:
                    # If the section itself is too large, split it by sentences
                    sentences = re.split(r'(?<=[.!?])\s+', section)
                    temp_chunk = ""
                    
                    for sentence in sentences:
                        if len(temp_chunk) + len(sentence) < max_chunk_size:
                            temp_chunk += sentence + " "
                        else:
                            if temp_chunk:
                                chunk_id += 1
                                chunk_metadata = {
                                    "chunk_id": f"{doc_id}_{chunk_id}",
                                    "chunk_index": chunk_id,
                                }
                                
                                if self.use_advanced_nlp:
                                    financial_entities = self._extract_financial_entities(temp_chunk)
                                    if financial_entities:
                                        chunk_metadata["financial_entities"] = financial_entities
                                
                                chunks.append(DocumentChunk(
                                    content=temp_chunk,
                                    metadata=chunk_metadata,
                                    chunk_id=f"{doc_id}_{chunk_id}"
                                ))
                            temp_chunk = sentence + " "
                    
                    if temp_chunk:
                        chunk_id += 1
                        chunk_metadata = {
                            "chunk_id": f"{doc_id}_{chunk_id}",
                            "chunk_index": chunk_id,
                        }
                        
                        if self.use_advanced_nlp:
                            financial_entities = self._extract_financial_entities(temp_chunk)
                            if financial_entities:
                                chunk_metadata["financial_entities"] = financial_entities
                        
                        chunks.append(DocumentChunk(
                            content=temp_chunk,
                            metadata=chunk_metadata,
                            chunk_id=f"{doc_id}_{chunk_id}"
                        ))
                    
                    current_chunk = ""
        
        # Add remaining content as a chunk
        if current_chunk:
            chunk_id += 1
            chunk_metadata = {
                "chunk_id": f"{doc_id}_{chunk_id}",
                "chunk_index": chunk_id,
            }
            
            if self.use_advanced_nlp:
                financial_entities = self._extract_financial_entities(current_chunk)
                if financial_entities:
                    chunk_metadata["financial_entities"] = financial_entities
            
            chunks.append(DocumentChunk(
                content=current_chunk,
                metadata=chunk_metadata,
                chunk_id=f"{doc_id}_{chunk_id}"
            ))
        
        return chunks
    
    def _chunk_tabular_document(self, content: str, metadata: Dict[str, Any], doc_id: str) -> List[DocumentChunk]:
        """
        Chunk a tabular document (Excel, CSV) by tables/sheets.
        """
        chunks = []
        
        # Split by sheet markers
        sheets = re.split(r'---\s*Sheet:\s*([^-]+)\s*---', content)
        
        if len(sheets) > 1:
            # Process each sheet as a separate chunk
            for i in range(1, len(sheets), 2):
                sheet_name = sheets[i]
                sheet_content = sheets[i+1] if i+1 < len(sheets) else ""
                
                chunk_id = f"{doc_id}_sheet_{i//2}"
                chunk_metadata = {
                    "chunk_id": chunk_id,
                    "chunk_index": i//2,
                    "sheet_name": sheet_name,
                    "content_type": "tabular",
                }
                
                chunks.append(DocumentChunk(
                    content=sheet_content,
                    metadata=chunk_metadata,
                    chunk_id=chunk_id
                ))
        else:
            # If no sheet markers, chunk by sections of the table
            lines = content.split('\n')
            
            # Assume header is in the first row
            header = lines[0] if lines else ""
            
            # Chunk the table into segments of related rows
            rows_per_chunk = 100
            current_chunk = header + "\n"
            chunk_count = 0
            
            for i, line in enumerate(lines[1:], 1):
                if i % rows_per_chunk == 0 and current_chunk:
                    chunk_count += 1
                    chunk_id = f"{doc_id}_table_{chunk_count}"
                    
                    chunks.append(DocumentChunk(
                        content=current_chunk,
                        metadata={
                            "chunk_id": chunk_id,
                            "chunk_index": chunk_count,
                            "content_type": "tabular",
                            "row_range": f"{i-rows_per_chunk+1}-{i}"
                        },
                        chunk_id=chunk_id
                    ))
                    
                    # Start new chunk with header
                    current_chunk = header + "\n" + line + "\n"
                else:
                    current_chunk += line + "\n"
            
            # Add the last chunk if it has content
            if current_chunk and current_chunk != header + "\n":
                chunk_count += 1
                chunk_id = f"{doc_id}_table_{chunk_count}"
                
                chunks.append(DocumentChunk(
                    content=current_chunk,
                    metadata={
                        "chunk_id": chunk_id,
                        "chunk_index": chunk_count,
                        "content_type": "tabular",
                        "row_range": f"{len(lines)-len(current_chunk.split('\\n'))+1}-{len(lines)}"
                    },
                    chunk_id=chunk_id
                ))
        
        return chunks
    
    def _basic_chunking(self, content: str, metadata: Dict[str, Any], doc_id: str) -> List[DocumentChunk]:
        """
        Basic chunking strategy that splits text into fixed-size chunks.
        Used as a fallback when more sophisticated strategies don't apply.
        """
        chunks = []
        chunk_size = 1000  # characters
        chunk_overlap = 200  # characters of overlap between chunks
        
        # Split content into chunks
        start = 0
        chunk_id = 0
        
        while start < len(content):
            # Create chunk with overlap
            end = min(start + chunk_size, len(content))
            
            # Try to end at a sentence boundary if possible
            if end < len(content):
                # Find the last sentence boundary in the chunk
                last_period = content.rfind('.', start, end)
                if last_period > start + chunk_size // 2:  # Only use if it's not too far back
                    end = last_period + 1
            
            chunk_text = content[start:end]
            chunk_id += 1
            
            chunk_metadata = {
                "chunk_id": f"{doc_id}_{chunk_id}",
                "chunk_index": chunk_id,
                "char_start": start,
                "char_end": end,
            }
            
            # Extract financial entities if advanced NLP is enabled
            if self.use_advanced_nlp:
                financial_entities = self._extract_financial_entities(chunk_text)
                if financial_entities:
                    chunk_metadata["financial_entities"] = financial_entities
            
            chunks.append(DocumentChunk(
                content=chunk_text,
                metadata=chunk_metadata,
                chunk_id=f"{doc_id}_{chunk_id}"
            ))
            
            # Move to next chunk with overlap
            start = end - chunk_overlap
        
        return chunks
    
    def _extract_financial_entities(self, text: str) -> Dict[str, List[str]]:
        """
        Extract financial entities from text using NLP.
        
        Returns a dictionary of entity types and their mentions.
        """
        entities = {
            "companies": [],
            "financial_terms": [],
            "monetary_values": [],
            "percentages": [],
            "dates": []
        }
        
        try:
            # Extract entities using spaCy
            doc = self.nlp(text)
            
            for ent in doc.ents:
                if ent.label_ == "ORG":
                    entities["companies"].append(ent.text)
                elif ent.label_ == "MONEY":
                    entities["monetary_values"].append(ent.text)
                elif ent.label_ == "PERCENT":
                    entities["percentages"].append(ent.text)
                elif ent.label_ == "DATE":
                    entities["dates"].append(ent.text)
            
            # Use financial entity extraction model
            financial_entities = self.financial_ner(text)
            
            for entity in financial_entities:
                # Map standard NER tags to our financial categories
                if entity["entity_group"] == "ORG":
                    entities["companies"].append(entity["word"])
                elif entity["entity_group"] == "PER":
                    # Sometimes key executives are mentioned
                    if "key_people" not in entities:
                        entities["key_people"] = []
                    entities["key_people"].append(entity["word"])
                elif entity["entity_group"] == "LOC":
                    # Locations can be relevant for jurisdictional analysis
                    if "locations" not in entities:
                        entities["locations"] = []
                    entities["locations"].append(entity["word"])
                else:
                    # Other entities might be financial terms
                    entities["financial_terms"].append(entity["word"])
                    
            # If we have the financial classifier, use it to determine sentiment
            if hasattr(self, 'financial_classifier') and self.financial_classifier:
                try:
                    sentiment_results = self.financial_classifier(text[:512])  # Limit text length for classification
                    if sentiment_results:
                        # Get the sentiment with highest score
                        sentiment = max(sentiment_results[0], key=lambda x: x['score'])
                        if "financial_sentiment" not in entities:
                            entities["financial_sentiment"] = []
                        entities["financial_sentiment"].append(f"{sentiment['label']}: {sentiment['score']:.2f}")
                except Exception as e:
                    logger.warning(f"Error in financial sentiment analysis: {e}")
            
            # Remove duplicates
            for entity_type in entities:
                entities[entity_type] = list(set(entities[entity_type]))
            
            # Only return non-empty entity types
            return {k: v for k, v in entities.items() if v}
            
        except Exception as e:
            logger.warning(f"Error extracting financial entities: {e}")
            return {}
    
    def extract_metadata(self, document: Document) -> Dict[str, Any]:
        """
        Extract metadata from a document, with a focus on financial information.
        
        Args:
            document: The document to extract metadata from.
            
        Returns:
            Extracted metadata.
        """
        metadata = {}
        content = document.content
        
        # Extract document type/category
        metadata["doc_category"] = self._categorize_financial_document(content)
        
        # Extract key financial information
        financial_info = self._extract_key_financial_info(content)
        if financial_info:
            metadata.update(financial_info)
        
        # Process using advanced NLP if enabled
        if self.use_advanced_nlp:
            # Extract all financial entities
            all_entities = self._extract_financial_entities(content)
            
            # Add to metadata
            for entity_type, entities in all_entities.items():
                if entities:
                    metadata[entity_type] = entities
            
            # Extract key financial metrics
            financial_metrics = self._extract_financial_metrics(content)
            if financial_metrics:
                metadata["financial_metrics"] = financial_metrics
        
        return metadata
    
    def _categorize_financial_document(self, content: str) -> str:
        """
        Categorize the financial document based on its content.
        
        Returns a category like: "financial_statement", "contract", "regulatory_filing", etc.
        """
        # Create category indicators - key terms associated with each document type
        categories = {
            "financial_statement": [
                "balance sheet", "income statement", "cash flow", "statement of financial position",
                "profit and loss", "annual report", "quarterly report", "10-K", "10-Q", "financial results"
            ],
            "due_diligence": [
                "due diligence", "DD report", "findings", "compliance assessment", 
                "risk assessment", "SWOT analysis"
            ],
            "contract": [
                "agreement", "contract", "terms and conditions", "parties agree", "hereby",
                "obligations", "termination clause", "effective date", "binding agreement"
            ],
            "regulatory_filing": [
                "SEC filing", "form 8-K", "proxy statement", "prospectus", "S-1", "registration statement",
                "disclosure", "regulatory compliance", "securities", "exchange commission"
            ],
            "valuation": [
                "valuation", "fair market value", "DCF", "discounted cash flow", "NPV",
                "net present value", "EBITDA multiple", "comparable analysis"
            ],
            "market_analysis": [
                "market analysis", "industry overview", "competitive landscape", "market share",
                "growth projections", "market size", "CAGR", "competitive advantage"
            ],
            "legal_memo": [
                "legal memorandum", "legal analysis", "legal opinion", "counsel", "liability",
                "legal risk", "jurisdiction", "statutory", "regulation", "compliance"
            ],
            "tax_document": [
                "tax", "taxation", "tax implications", "tax structure", "tax liability",
                "depreciation", "amortization", "tax-free", "tax benefit", "IRS"
            ]
        }
        
        # Count occurrences of each category's indicators
        category_scores = {category: 0 for category in categories}
        
        for category, indicators in categories.items():
            for indicator in indicators:
                # Count case-insensitive occurrences of the indicator
                count = len(re.findall(r'\b' + re.escape(indicator) + r'\b', content, re.IGNORECASE))
                category_scores[category] += count
        
        # Get the category with the highest score
        if any(category_scores.values()):
            best_category = max(category_scores.items(), key=lambda x: x[1])[0]
            return best_category
        
        # Default category if no clear match
        return "general_financial"
    
    def _extract_key_financial_info(self, content: str) -> Dict[str, Any]:
        """
        Extract key financial information from document content.
        """
        financial_info = {}
        
        # Extract potential company names
        company_pattern = r'(?:(?:Inc|LLC|Ltd|Corporation|Corp|Company|Co|LP|LLP|SA|GmbH|Plc|AG|NV)\.?)'
        company_names = re.findall(r'[A-Z][a-zA-Z0-9\s&,]+' + company_pattern, content)
        if company_names:
            financial_info["potential_companies"] = list(set(company_names[:10]))  # Limit to top 10
        
        # Extract monetary values
        money_pattern = r'\$\s*\d+(?:,\d{3})*(?:\.\d+)?(?:\s*(?:million|billion|m|b|M|B))?|\d+(?:,\d{3})*(?:\.\d+)?\s*(?:USD|EUR|GBP|JPY|CAD|AUD)'
        monetary_values = re.findall(money_pattern, content)
        if monetary_values:
            financial_info["monetary_values"] = list(set(monetary_values[:10]))  # Limit to top 10
        
        # Extract dates (common financial document date formats)
        date_pattern = r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}\b|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
        dates = re.findall(date_pattern, content)
        if dates:
            financial_info["dates"] = list(set(dates[:10]))  # Limit to top 10
        
        # Extract percentages
        percentage_pattern = r'\b\d+(?:\.\d+)?%\b'
        percentages = re.findall(percentage_pattern, content)
        if percentages:
            financial_info["percentages"] = list(set(percentages[:10]))  # Limit to top 10
        
        return financial_info
    
    def _extract_financial_metrics(self, content: str) -> Dict[str, Any]:
        """
        Extract key financial metrics from the document.
        """
        metrics = {}
        
        # Define patterns for common financial metrics
        metric_patterns = {
            "revenue": r'(?:revenue|sales|turnover)(?:s)?\s*(?:of|:)?\s*(?:US)?\$?\s*(\d+(?:,\d{3})*(?:\.\d+)?(?:\s*(?:million|billion|m|b|M|B))?)',
            "ebitda": r'EBITDA\s*(?:of|:)?\s*(?:US)?\$?\s*(\d+(?:,\d{3})*(?:\.\d+)?(?:\s*(?:million|billion|m|b|M|B))?)',
            "net_income": r'(?:net income|profit|earnings)(?:s)?\s*(?:of|:)?\s*(?:US)?\$?\s*(\d+(?:,\d{3})*(?:\.\d+)?(?:\s*(?:million|billion|m|b|M|B))?)',
            "assets": r'(?:total assets|assets)(?:s)?\s*(?:of|:)?\s*(?:US)?\$?\s*(\d+(?:,\d{3})*(?:\.\d+)?(?:\s*(?:million|billion|m|b|M|B))?)',
            "liabilities": r'(?:total liabilities|liabilities)(?:s)?\s*(?:of|:)?\s*(?:US)?\$?\s*(\d+(?:,\d{3})*(?:\.\d+)?(?:\s*(?:million|billion|m|b|M|B))?)',
            "equity": r'(?:equity|shareholder\'s equity|stockholder\'s equity)(?:s)?\s*(?:of|:)?\s*(?:US)?\$?\s*(\d+(?:,\d{3})*(?:\.\d+)?(?:\s*(?:million|billion|m|b|M|B))?)',
            "pe_ratio": r'(?:price[- ]to[- ]earnings|P/E|PE) ratio\s*(?:of|:)?\s*(\d+(?:\.\d+)?)',
            "market_cap": r'(?:market capitalization|market cap)(?:s)?\s*(?:of|:)?\s*(?:US)?\$?\s*(\d+(?:,\d{3})*(?:\.\d+)?(?:\s*(?:million|billion|m|b|M|B))?)'
        }
        
        # Search for each metric in the content
        for metric_name, pattern in metric_patterns.items():
            matches = re.findall(pattern, content, re.IGNORECASE)
            if matches:
                metrics[metric_name] = matches[0]  # Take the first match
        
        return metrics
